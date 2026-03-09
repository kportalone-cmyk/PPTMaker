"""
DOCX 생성 서비스 - python-docx 기반 Word 문서 생성

generated_docx 컬렉션의 데이터를 .docx 파일로 변환합니다.
웹 프리뷰(CKEditor/스트리밍)와 최대한 동일한 스타일을 적용합니다.
"""

import os
import re
import json
import uuid
import tempfile
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from services.mongo_service import get_db
from config import settings

# matplotlib 차트 렌더링
try:
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    import numpy as np
    _HAS_MATPLOTLIB = True
except ImportError:
    _HAS_MATPLOTLIB = False

# 웹 프리뷰와 동일한 색상 팔레트
_COLORS = {
    "h1": RGBColor(0x1A, 0x1A, 0x2E),
    "h2": RGBColor(0x1E, 0x29, 0x3B),
    "h3": RGBColor(0x33, 0x41, 0x55),
    "h4": RGBColor(0x47, 0x55, 0x69),
    "body": RGBColor(0x33, 0x41, 0x55),
    "description": RGBColor(0x64, 0x74, 0x8B),
    "blockquote": RGBColor(0x43, 0x38, 0xCA),
    "table_header_text": RGBColor(0x1E, 0x29, 0x3B),
    "table_body_text": RGBColor(0x33, 0x41, 0x55),
    "accent": RGBColor(0x63, 0x66, 0xF1),
}

# 제목 폰트 크기 (px→pt 변환: CKEditor 14px 기준)
_HEADING_SIZES = {
    0: Pt(22),  # 문서 타이틀
    1: Pt(20),  # h1: 26px → ~20pt
    2: Pt(15),  # h2: 20px → ~15pt
    3: Pt(13),  # h3: 17px → ~13pt
    4: Pt(11),  # h4: 15px → ~11pt
}

# 제목 색상
_HEADING_COLORS = {
    0: _COLORS["h1"],
    1: _COLORS["h1"],
    2: _COLORS["h2"],
    3: _COLORS["h3"],
    4: _COLORS["h4"],
}


def _set_cell_shading(cell, color_hex: str):
    """테이블 셀 배경색 설정"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _set_paragraph_spacing(paragraph, line_spacing=1.6, space_before=0, space_after=Pt(6)):
    """단락 줄간격 및 전후 간격 설정"""
    pf = paragraph.paragraph_format
    pf.line_spacing = line_spacing
    if space_before is not None:
        pf.space_before = space_before
    if space_after is not None:
        pf.space_after = space_after


def _set_run_font(run, font_name="맑은 고딕", size=None, color=None, bold=None, italic=None):
    """run에 폰트 속성 설정"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if size:
        run.font.size = size
    if color:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def _add_styled_heading(doc, text: str, level: int):
    """웹 프리뷰 스타일과 일치하는 제목 추가"""
    level = min(level, 4)
    p = doc.add_paragraph()

    # 제목 텍스트 run
    run = p.add_run(text)
    _set_run_font(
        run,
        size=_HEADING_SIZES.get(level, Pt(11)),
        color=_HEADING_COLORS.get(level, _COLORS["body"]),
        bold=True,
    )

    # 줄간격 설정
    pf = p.paragraph_format
    pf.line_spacing = 1.4

    if level == 0:
        # 문서 타이틀: 가운데 정렬, 아래 여백
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf.space_before = Pt(4)
        pf.space_after = Pt(20)
        # 하단 보더 (accent 색상)
        p_element = p._element
        pPr = p_element.get_or_add_pPr()
        borders = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'  <w:bottom w:val="single" w:sz="12" w:space="4" w:color="6366F1"/>'
            f'</w:pBdr>'
        )
        pPr.append(borders)

    elif level == 1:
        # h1: 큰 여백
        pf.space_before = Pt(36)
        pf.space_after = Pt(14)
        # 하단 보더
        p_element = p._element
        pPr = p_element.get_or_add_pPr()
        borders = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'  <w:bottom w:val="single" w:sz="12" w:space="4" w:color="6366F1"/>'
            f'</w:pBdr>'
        )
        pPr.append(borders)

    elif level == 2:
        # h2: 좌측 보더 + 배경
        pf.space_before = Pt(30)
        pf.space_after = Pt(12)
        pf.left_indent = Pt(10)
        # 좌측 보더
        p_element = p._element
        pPr = p_element.get_or_add_pPr()
        borders = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'  <w:left w:val="single" w:sz="18" w:space="8" w:color="6366F1"/>'
            f'</w:pBdr>'
        )
        pPr.append(borders)
        # 배경 shading
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0F0FF" w:val="clear"/>')
        pPr.append(shading)

    elif level == 3:
        # h3: 얇은 좌측 보더
        pf.space_before = Pt(24)
        pf.space_after = Pt(10)
        pf.left_indent = Pt(8)
        p_element = p._element
        pPr = p_element.get_or_add_pPr()
        borders = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'  <w:left w:val="single" w:sz="12" w:space="6" w:color="A5B4FC"/>'
            f'</w:pBdr>'
        )
        pPr.append(borders)

    elif level == 4:
        # h4: 간단한 여백만
        pf.space_before = Pt(20)
        pf.space_after = Pt(8)

    return p


# ── 차트 렌더링 ──────────────────────────────────────────

# 차트 색상 팔레트 (웹 프리뷰와 동일)
_CHART_PALETTE = [
    "#6366F1", "#EC4899", "#F59E0B", "#10B981",
    "#3B82F6", "#8B5CF6", "#EF4444", "#14B8A6",
]


def _is_chart_json(data: dict) -> bool:
    """dict가 Chart.js 설정인지 확인"""
    return (
        isinstance(data, dict)
        and "type" in data
        and "data" in data
        and isinstance(data.get("data"), dict)
        and "labels" in data["data"]
        and "datasets" in data["data"]
    )


def _render_chart_image(chart_config: dict, output_path: str) -> bool:
    """Chart.js JSON 설정을 matplotlib 이미지로 변환"""
    if not _HAS_MATPLOTLIB:
        return False

    try:
        # 한글 폰트 설정
        plt.rcParams['font.family'] = 'Malgun Gothic'
        plt.rcParams['axes.unicode_minus'] = False

        chart_type = chart_config.get("type", "bar")
        data = chart_config["data"]
        labels = data.get("labels", [])
        datasets = data.get("datasets", [])
        options = chart_config.get("options", {})

        if not labels or not datasets:
            return False

        fig, ax = plt.subplots(figsize=(10, 5.5))

        if chart_type == 'bar':
            x = np.arange(len(labels))
            n = len(datasets)
            width = 0.7 / max(n, 1)
            for idx, ds in enumerate(datasets):
                offset = (idx - n / 2 + 0.5) * width
                colors = ds.get("backgroundColor", _CHART_PALETTE[idx % len(_CHART_PALETTE)])
                ax.bar(x + offset, ds["data"], width,
                       label=ds.get("label", ""), color=colors,
                       edgecolor='white', linewidth=0.5)
            ax.set_xticks(x)
            max_label_len = max((len(str(l)) for l in labels), default=0)
            rotation = 30 if max_label_len > 6 else 0
            ha = 'right' if rotation else 'center'
            ax.set_xticklabels(labels, rotation=rotation, ha=ha, fontsize=10)

        elif chart_type == 'line':
            for idx, ds in enumerate(datasets):
                color = ds.get("borderColor", _CHART_PALETTE[idx % len(_CHART_PALETTE)])
                ax.plot(range(len(labels)), ds["data"],
                        label=ds.get("label", ""), marker='o',
                        linewidth=2.5, color=color, markersize=6)
            ax.set_xticks(range(len(labels)))
            ax.set_xticklabels(labels, fontsize=10)

        elif chart_type in ('pie', 'doughnut'):
            ds = datasets[0]
            colors = ds.get("backgroundColor", _CHART_PALETTE[:len(labels)])
            if chart_type == 'doughnut':
                wedges, texts, autotexts = ax.pie(
                    ds["data"], labels=labels, autopct='%1.1f%%',
                    colors=colors, pctdistance=0.75,
                    textprops={'fontsize': 10})
                centre = plt.Circle((0, 0), 0.50, fc='white')
                ax.add_patch(centre)
            else:
                ax.pie(ds["data"], labels=labels, autopct='%1.1f%%',
                       colors=colors, textprops={'fontsize': 10})
            ax.set_aspect('equal')

        elif chart_type == 'area':
            for idx, ds in enumerate(datasets):
                color = ds.get("backgroundColor", _CHART_PALETTE[idx % len(_CHART_PALETTE)])
                ax.fill_between(range(len(labels)), ds["data"],
                                alpha=0.4, color=color, label=ds.get("label", ""))
                ax.plot(range(len(labels)), ds["data"], linewidth=2, color=color)
            ax.set_xticks(range(len(labels)))
            ax.set_xticklabels(labels, fontsize=10)

        else:
            # 기본 bar
            x = np.arange(len(labels))
            for idx, ds in enumerate(datasets):
                colors = ds.get("backgroundColor", _CHART_PALETTE[idx % len(_CHART_PALETTE)])
                ax.bar(x, ds["data"], label=ds.get("label", ""), color=colors)
            ax.set_xticks(x)
            ax.set_xticklabels(labels, fontsize=10)

        # 제목
        title = options.get("title", "")
        if isinstance(title, str) and title:
            ax.set_title(title, fontsize=14, fontweight='bold', pad=15)
        elif isinstance(title, dict):
            ax.set_title(title.get("text", title.get("display", "")),
                         fontsize=14, fontweight='bold', pad=15)

        # 축 레이블
        scales = options.get("scales", {})
        if isinstance(scales, dict):
            y_cfg = scales.get("y", {})
            if isinstance(y_cfg, dict):
                y_title = y_cfg.get("title", "")
                if isinstance(y_title, str) and y_title:
                    ax.set_ylabel(y_title, fontsize=11)
                elif isinstance(y_title, dict):
                    ax.set_ylabel(y_title.get("text", ""), fontsize=11)
            x_cfg = scales.get("x", {})
            if isinstance(x_cfg, dict):
                x_title = x_cfg.get("title", "")
                if isinstance(x_title, str) and x_title:
                    ax.set_xlabel(x_title, fontsize=11)
                elif isinstance(x_title, dict):
                    ax.set_xlabel(x_title.get("text", ""), fontsize=11)

        # 범례 및 그리드
        if chart_type not in ('pie', 'doughnut'):
            if any(ds.get("label") for ds in datasets):
                ax.legend(fontsize=10, framealpha=0.9)
            ax.grid(axis='y', alpha=0.3, linestyle='--')
            ax.spines['top'].set_visible(False)
            ax.spines['right'].set_visible(False)

        plt.tight_layout()
        fig.savefig(output_path, dpi=150, bbox_inches='tight', facecolor='white')
        plt.close(fig)
        return True

    except Exception as e:
        print(f"[WordService] Chart render error: {e}")
        try:
            plt.close('all')
        except:
            pass
        return False


def _add_chart_to_doc(doc: Document, chart_config: dict):
    """Chart.js 설정을 이미지로 렌더링하여 문서에 삽입"""
    tmp_path = os.path.join(tempfile.gettempdir(), f"chart_{uuid.uuid4().hex}.png")
    try:
        if _render_chart_image(chart_config, tmp_path):
            doc.add_picture(tmp_path, width=Cm(15))
            last_p = doc.paragraphs[-1]
            last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_paragraph_spacing(last_p, space_before=Pt(10), space_after=Pt(14))
    finally:
        try:
            if os.path.exists(tmp_path):
                os.remove(tmp_path)
        except:
            pass


def _add_code_block(doc: Document, code_text: str):
    """코드 블록을 배경색 있는 단락으로 렌더링"""
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    _set_run_font(run, font_name="Consolas", size=Pt(9), color=_COLORS["body"])
    _set_paragraph_spacing(p, line_spacing=1.3, space_before=Pt(6), space_after=Pt(6))
    # 배경색
    p_element = p._element
    pPr = p_element.get_or_add_pPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F1F5F9" w:val="clear"/>')
    pPr.append(shading)


def extract_docx_template_info(file_path: str) -> dict:
    """업로드된 .docx 템플릿에서 스타일 정보 및 콘텐츠 구조 추출"""
    try:
        doc = Document(file_path)
        info = {
            "page_count": len(doc.sections),
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
            "styles": [],
            "structure": [],  # 템플릿 섹션 구조
        }
        # 주요 스타일 추출
        seen = set()
        for para in doc.paragraphs:
            style_name = para.style.name
            if style_name not in seen:
                seen.add(style_name)
                font = para.style.font
                info["styles"].append({
                    "name": style_name,
                    "font": font.name if font.name else None,
                    "size": font.size.pt if font.size else None,
                    "bold": font.bold,
                })

        # 콘텐츠 구조 추출 - 섹션 제목과 안내 텍스트 파악
        info["structure"] = extract_docx_template_structure(file_path)

        return info
    except Exception as e:
        print(f"[WordService] 템플릿 분석 실패: {e}")
        return {}


def _looks_like_placeholder(text: str) -> bool:
    """텍스트가 플레이스홀더/안내 텍스트인지 판별"""
    text = text.strip()
    if not text or len(text) > 200:
        return False
    patterns = [
        '입력해', '작성해', '기입해', '기재해', '기술해',
        '입력하세요', '작성하세요', '기입하세요',
        '주세요', '바랍니다', '하십시오',
        '여기에', '내용을', '입력란',
    ]
    return any(p in text for p in patterns)


def _extract_table_based_structure(doc: Document) -> list:
    """테이블 레이아웃 기반 템플릿에서 섹션 구조 추출

    한국어 워드 템플릿은 테이블로 레이아웃을 구성하고,
    섹션 헤더를 이미지로, 본문을 셀 텍스트로 배치하는 경우가 많습니다.
    병합된 셀과 복잡한 테이블 구조를 안전하게 처리합니다.
    """
    structure = []
    seen_texts = set()  # (table_index, cell_text) 쌍으로 중복 방지

    for ti, table in enumerate(doc.tables):
        try:
            for ri, row in enumerate(table.rows):
                try:
                    cells = row.cells
                except Exception:
                    continue
                for ci, cell in enumerate(cells):
                    try:
                        cell_text = cell.text.strip()
                        if not cell_text:
                            continue

                        # 같은 테이블 내 동일 텍스트 중복 방지 (병합 셀)
                        dedup_key = (ti, cell_text)
                        if dedup_key in seen_texts:
                            continue
                        seen_texts.add(dedup_key)

                        if _looks_like_placeholder(cell_text):
                            structure.append({
                                "title": cell_text,
                                "level": 1,
                                "placeholder": cell_text,
                                "type": "table_cell",
                                "table_index": ti,
                                "cell_row": ri,
                                "cell_col": ci,
                            })
                    except Exception:
                        continue
        except Exception as e:
            print(f"[WordService] 테이블 {ti} 분석 중 에러: {e}")
            continue

    return structure


def extract_docx_template_structure(file_path: str) -> list:
    """템플릿 .docx에서 섹션 구조(제목 + 안내/플레이스홀더 텍스트) 추출

    단락 기반 템플릿과 테이블 기반 템플릿 모두 지원합니다.

    Returns:
        [
            {"title": "섹션 제목", "level": 1, "placeholder": "안내 텍스트 또는 빈 문자열"},
            ...
        ]
    """
    try:
        doc = Document(file_path)

        # 1단계: 단락 기반 구조 추출
        structure = []
        current_section = None

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            if _is_template_heading(para):
                if current_section:
                    structure.append(current_section)
                # 레벨 결정
                style_name = para.style.name.lower()
                level = 1
                if 'title' in style_name:
                    level = 0
                elif 'heading' in style_name:
                    for i in range(1, 5):
                        if str(i) in style_name:
                            level = i
                            break
                else:
                    runs = para.runs
                    if runs:
                        font_size = runs[0].font.size or (para.style.font.size if para.style.font else None)
                        if font_size:
                            if font_size >= Pt(20):
                                level = 0
                            elif font_size >= Pt(16):
                                level = 1
                            else:
                                level = 2
                current_section = {"title": text, "level": level, "placeholder": ""}
            elif current_section:
                if current_section["placeholder"]:
                    current_section["placeholder"] += "\n" + text
                else:
                    current_section["placeholder"] = text

        if current_section:
            structure.append(current_section)

        if structure:
            print(f"[WordService] 단락 기반 템플릿 구조: {len(structure)}개 섹션")
            return structure

        # 2단계: 단락에서 못 찾으면 테이블 기반 구조 추출
        table_structure = _extract_table_based_structure(doc)
        if table_structure:
            print(f"[WordService] 테이블 기반 템플릿 구조: {len(table_structure)}개 섹션")
            return table_structure

        print("[WordService] 템플릿 구조 추출 실패: 단락/테이블 모두 감지 안됨")
        return []
    except Exception as e:
        print(f"[WordService] 템플릿 구조 추출 실패: {e}")
        return []


def _load_template_document(template_path: str) -> Document:
    """템플릿 .docx를 로드하고 기존 본문 내용을 비운 뒤 반환 (스타일/헤더/푸터 유지)"""
    doc = Document(template_path)
    # 본문 콘텐츠 제거 (스타일, 섹션 설정, 머리글/바닥글 유지)
    body = doc.element.body
    # 모든 단락과 테이블 제거 (sectPr은 유지)
    for child in list(body):
        if child.tag.endswith('}sectPr'):
            continue
        body.remove(child)
    return doc


# ── 템플릿 기반 생성 헬퍼 ──────────────────────────────────


def _is_template_heading(para) -> bool:
    """단락이 템플릿의 제목/헤딩인지 판별"""
    text = para.text.strip()
    if not text:
        return False

    style_name = para.style.name.lower()
    if 'heading' in style_name or 'title' in style_name:
        return True

    # 스타일이 Heading이 아니어도 굵은 글씨 + 큰 폰트이면 제목으로 간주
    runs = para.runs
    if runs:
        first_run = runs[0]
        is_bold = first_run.bold or (para.style.font.bold if para.style.font else False)
        font_size = first_run.font.size or (para.style.font.size if para.style.font else None)
        if is_bold and font_size and font_size >= Pt(13):
            return True

    return False


def _identify_template_sections(doc: Document) -> list:
    """템플릿 문서에서 섹션(헤딩 + 플레이스홀더 요소) 식별

    Returns:
        [{"title": str, "heading_element": Element, "placeholder_elements": [Element]}]
    """
    sections = []
    current = None

    for para in doc.paragraphs:
        text = para.text.strip()

        if _is_template_heading(para):
            if current:
                sections.append(current)
            current = {
                'title': text,
                'heading_element': para._element,
                'placeholder_elements': [],
            }
        elif current and text:
            # 텍스트가 있는 단락만 플레이스홀더로 간주 (빈 단락/장식 요소 보존)
            current['placeholder_elements'].append(para._element)

    if current:
        sections.append(current)

    return sections


def _find_matching_ai_section(heading_title: str, ai_map: dict) -> dict:
    """템플릿 헤딩 제목으로 AI 섹션 매칭 (정확 → 번호제거 → 부분 매칭)"""
    # 정확 매칭
    if heading_title in ai_map:
        return ai_map.pop(heading_title)

    # 번호 제거 후 매칭
    clean_title = re.sub(r'^\d+[\.\)]\s*', '', heading_title).strip()
    for key in list(ai_map.keys()):
        clean_key = re.sub(r'^\d+[\.\)]\s*', '', key).strip()
        if clean_title == clean_key:
            return ai_map.pop(key)

    # 부분 포함 매칭
    for key in list(ai_map.keys()):
        clean_key = re.sub(r'^\d+[\.\)]\s*', '', key).strip()
        if clean_title and clean_key and len(clean_title) >= 2 and len(clean_key) >= 2:
            if clean_title in clean_key or clean_key in clean_title:
                return ai_map.pop(key)

    return None


def _insert_content_after(doc: Document, after_element, markdown_content: str):
    """마크다운 콘텐츠를 렌더링하여 특정 요소 뒤에 삽입"""
    body = doc.element.body

    # 렌더링 전 기존 요소를 lxml 요소 set으로 보관 (id() 대신 lxml 네이티브 비교)
    existing = set(body)

    # 콘텐츠를 문서 끝에 렌더링
    _render_markdown_content(doc, markdown_content)

    # 새로 추가된 요소 식별 (sectPr 제외)
    new_elements = [child for child in list(body)
                    if child not in existing and not child.tag.endswith('}sectPr')]

    if not new_elements:
        return

    # 새 요소들을 body에서 제거
    for elem in new_elements:
        body.remove(elem)

    # after_element 뒤에 순서대로 삽입
    insert_point = after_element
    for elem in new_elements:
        insert_point.addnext(elem)
        insert_point = elem


def _render_content_to_cell(doc: Document, cell, markdown_content: str):
    """마크다운 콘텐츠를 렌더링하여 테이블 셀에 삽입"""
    body = doc.element.body
    tc = cell._tc

    # 렌더링 전 기존 요소를 lxml 요소 set으로 보관 (id() 대신 lxml 네이티브 비교)
    existing = set(body)

    # 문서 끝에 렌더링
    _render_markdown_content(doc, markdown_content)

    # 새 요소 수집 (lxml 네이티브 비교로 안전하게 식별)
    new_elements = [child for child in list(body)
                    if child not in existing and not child.tag.endswith('}sectPr')]

    # body에서 제거
    for elem in new_elements:
        body.remove(elem)

    # 셀의 기존 내용 제거
    for p_elem in list(tc.findall(qn('w:p'))):
        tc.remove(p_elem)
    for t_elem in list(tc.findall(qn('w:tbl'))):
        tc.remove(t_elem)

    # 새 요소를 셀에 추가
    if new_elements:
        for elem in new_elements:
            tc.append(elem)
    else:
        # Word 규약: 셀에 최소 하나의 단락 필요
        tc.append(parse_xml(f'<w:p {nsdecls("w")}/>'))


def _fill_table_template(doc: Document, sections: list, template_structure: list):
    """테이블 기반 템플릿의 플레이스홀더 셀을 AI 콘텐츠로 교체"""
    # AI 섹션을 순서대로 매핑 (테이블 기반은 순서 매칭)
    ai_map = {}
    for s in sections:
        title = s.get('title', '').strip()
        if title:
            ai_map[title] = s

    matched = 0
    for ts in template_structure:
        if ts.get('type') != 'table_cell':
            continue

        ti = ts['table_index']
        ri = ts['cell_row']
        ci = ts['cell_col']
        placeholder = ts.get('placeholder', '')

        if ti >= len(doc.tables):
            continue

        table = doc.tables[ti]
        if ri >= len(table.rows):
            continue

        try:
            row_cells = table.rows[ri].cells
            if ci >= len(row_cells):
                continue
            cell = row_cells[ci]
        except Exception:
            continue

        # AI 섹션 매칭 (제목=플레이스홀더 텍스트)
        ai_section = _find_matching_ai_section(placeholder, ai_map)

        if not ai_section:
            # 순서 기반 폴백 매칭
            if matched < len(sections):
                ai_section = sections[matched]
            else:
                print(f"[WordService] 테이블 셀 매칭 실패: [{ti},{ri},{ci}] '{placeholder[:40]}'")
                continue

        content = ai_section.get('content', '')
        if not content:
            matched += 1
            continue

        print(f"[WordService] 테이블 셀 매칭: [{ti},{ri},{ci}] → 콘텐츠 {len(content)}자")

        # 셀 내용 교체
        _render_content_to_cell(doc, cell, content)
        matched += 1


def _generate_from_template(template_path: str, sections: list, meta: dict) -> Document:
    """템플릿의 구조/서식을 보존하면서 AI 콘텐츠를 채워넣기"""
    doc = Document(template_path)
    body = doc.element.body

    print(f"[WordService] 템플릿 기반 문서 생성: {template_path}")

    # 1단계: 단락 기반 섹션 식별
    template_sections = _identify_template_sections(doc)

    if template_sections:
        # 단락 기반 템플릿 처리
        ai_map = {}
        for s in sections:
            title = s.get('title', '').strip()
            if title:
                ai_map[title] = s

        print(f"[WordService] 단락 기반 템플릿: {len(template_sections)}개, AI: {len(ai_map)}개")

        for ts in reversed(template_sections):
            heading_title = ts['title']
            placeholder_elements = ts['placeholder_elements']
            heading_element = ts['heading_element']

            ai_section = _find_matching_ai_section(heading_title, ai_map)
            if not ai_section:
                print(f"[WordService] 매칭 실패: '{heading_title}'")
                continue

            content = ai_section.get('content', '')
            if not content:
                continue

            print(f"[WordService] 섹션 매칭: '{heading_title}' → {len(content)}자")

            for elem in placeholder_elements:
                body.remove(elem)

            _insert_content_after(doc, heading_element, content)

        return doc

    # 2단계: 테이블 기반 섹션 식별
    table_structure = _extract_table_based_structure(doc)

    if table_structure:
        print(f"[WordService] 테이블 기반 템플릿: {len(table_structure)}개 섹션")
        _fill_table_template(doc, sections, table_structure)
        return doc

    # 3단계: 모두 실패 → 기존 방식 폴백
    print("[WordService] 템플릿 섹션 식별 실패 → 기본 방식 폴백")
    doc = _load_template_document(template_path)
    _fill_document_content(doc, sections, meta)
    return doc


def _fill_document_content(doc: Document, sections: list, meta: dict):
    """문서에 섹션 콘텐츠 채우기 (기본 하드코딩 스타일)"""
    # 문서 제목
    title = meta.get("title", "")
    if title:
        _add_styled_heading(doc, title, level=0)

    # 문서 설명
    description = meta.get("description", "")
    if description:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(description)
        _set_run_font(run, color=_COLORS["description"], size=Pt(10), italic=True)
        _set_paragraph_spacing(p, space_before=0, space_after=Pt(4))
        hr = doc.add_paragraph()
        hr_pf = hr.paragraph_format
        hr_pf.space_before = Pt(8)
        hr_pf.space_after = Pt(12)
        hr_element = hr._element
        hrPr = hr_element.get_or_add_pPr()
        hr_border = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="CBD5E1"/>'
            f'</w:pBdr>'
        )
        hrPr.append(hr_border)

    # 섹션 처리
    for section in sections:
        section_title = section.get("title", "")
        level = section.get("level", 1)
        content = section.get("content", "")

        if section_title:
            _add_styled_heading(doc, section_title, level=min(level, 4))

        if content:
            _render_markdown_content(doc, content)


async def create_empty_docx(project_id: str, template_path: str = None) -> str:
    """빈 docx 파일 생성 (OnlyOffice 에디터용) - 템플릿 지원"""
    if template_path and os.path.exists(template_path):
        doc = _load_template_document(template_path)
    else:
        doc = Document()
        # 기본 스타일 설정
        style = doc.styles["Normal"]
        font = style.font
        font.name = "맑은 고딕"
        font.size = Pt(11)
        font.color.rgb = _COLORS["body"]
        style.element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        pf = style.paragraph_format
        pf.line_spacing = 1.6
        pf.space_after = Pt(6)

    output_dir = os.path.join(settings.UPLOAD_DIR, "documents")
    os.makedirs(output_dir, exist_ok=True)
    filename = f"{uuid.uuid4().hex}.docx"
    output_path = os.path.join(output_dir, filename)
    doc.save(output_path)

    return f"/uploads/documents/{filename}"


async def generate_docx(project_id: str, template_path: str = None) -> str:
    """generated_docx 데이터를 .docx 파일로 변환"""
    db = get_db()
    docx_data = await db.generated_docx.find_one({"project_id": project_id})
    if not docx_data:
        raise ValueError("생성된 문서 데이터가 없습니다")

    sections = docx_data.get("sections", [])
    if not sections:
        raise ValueError("섹션 데이터가 없습니다")

    meta = docx_data.get("meta", {})

    # 템플릿 기반 생성: 템플릿 구조/서식 보존 + 플레이스홀더를 AI 콘텐츠로 교체
    if template_path and os.path.exists(template_path):
        doc = _generate_from_template(template_path, sections, meta)
    else:
        # 기본 생성 (하드코딩 스타일)
        doc = Document()
        style = doc.styles["Normal"]
        font = style.font
        font.name = "맑은 고딕"
        font.size = Pt(11)
        font.color.rgb = _COLORS["body"]
        style.element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        pf = style.paragraph_format
        pf.line_spacing = 1.6
        pf.space_after = Pt(6)
        _fill_document_content(doc, sections, meta)

    # 파일 저장
    output_dir = os.path.join(settings.UPLOAD_DIR, "documents")
    os.makedirs(output_dir, exist_ok=True)
    filename = f"{uuid.uuid4().hex}.docx"
    output_path = os.path.join(output_dir, filename)
    doc.save(output_path)

    return f"/uploads/documents/{filename}"


def _render_markdown_content(doc: Document, content: str):
    """마크다운 텍스트를 docx 요소로 변환"""
    lines = content.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 빈 줄
        if not stripped:
            i += 1
            continue

        # 코드 블록 (```...```) - 차트 JSON 감지 포함
        if stripped.startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1  # skip closing ```

            code_text = "\n".join(code_lines).strip()
            if code_text:
                # Chart.js JSON인지 확인
                try:
                    parsed = json.loads(code_text)
                    if _is_chart_json(parsed):
                        _add_chart_to_doc(doc, parsed)
                        continue
                except (json.JSONDecodeError, ValueError):
                    pass
                # 일반 코드 블록 렌더링
                _add_code_block(doc, code_text)
            continue

        # Raw JSON 차트 감지 (``` 없이 { 로 시작하는 JSON)
        if stripped == "{":
            json_lines = [lines[i]]
            brace_count = stripped.count("{") - stripped.count("}")
            j = i + 1
            while j < len(lines) and brace_count > 0:
                json_lines.append(lines[j])
                brace_count += lines[j].count("{") - lines[j].count("}")
                j += 1
            json_text = "\n".join(json_lines).strip()
            try:
                parsed = json.loads(json_text)
                if _is_chart_json(parsed):
                    _add_chart_to_doc(doc, parsed)
                    i = j
                    continue
            except (json.JSONDecodeError, ValueError):
                pass
            # 차트가 아니면 일반 텍스트로 처리 (fall through)

        # 마크다운 헤딩 (# ~ ####)
        heading_match = re.match(r"^(#{1,4})\s+(.+)$", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            heading_text = heading_match.group(2)
            _add_styled_heading(doc, heading_text, level=level)
            i += 1
            continue

        # 테이블 감지 (| col1 | col2 |)
        if stripped.startswith("|") and "|" in stripped[1:]:
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            _add_table(doc, table_lines)
            continue

        # 불릿 리스트
        if stripped.startswith("- ") or stripped.startswith("* "):
            text = stripped[2:]
            try:
                p = doc.add_paragraph(style="List Bullet")
            except KeyError:
                p = doc.add_paragraph()
                p.add_run("• ")
            _add_inline_formatting(p, text)
            _set_paragraph_spacing(p, space_before=Pt(2), space_after=Pt(2))
            i += 1
            continue

        # 번호 리스트
        num_match = re.match(r"^\d+\.\s+(.+)$", stripped)
        if num_match:
            text = num_match.group(1)
            try:
                p = doc.add_paragraph(style="List Number")
            except KeyError:
                p = doc.add_paragraph()
            _add_inline_formatting(p, text)
            _set_paragraph_spacing(p, space_before=Pt(2), space_after=Pt(2))
            i += 1
            continue

        # 인용문 (웹 프리뷰 인디고 테마와 동일)
        if stripped.startswith("> "):
            text = stripped[2:]
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.left_indent = Cm(1)
            _set_paragraph_spacing(p, line_spacing=1.5, space_before=Pt(6), space_after=Pt(6))
            # 좌측 보더 (accent 색상)
            p_element = p._element
            pPr = p_element.get_or_add_pPr()
            borders = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'  <w:left w:val="single" w:sz="18" w:space="8" w:color="6366F1"/>'
                f'</w:pBdr>'
            )
            pPr.append(borders)
            # 배경 shading
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0F0FF" w:val="clear"/>')
            pPr.append(shading)
            # 텍스트
            run = p.add_run(text)
            _set_run_font(run, color=_COLORS["blockquote"], size=Pt(10))
            i += 1
            continue

        # 일반 텍스트
        p = doc.add_paragraph()
        _add_inline_formatting(p, stripped)
        _set_paragraph_spacing(p, space_before=Pt(3), space_after=Pt(6))
        i += 1


def _add_inline_formatting(paragraph, text: str):
    """인라인 마크다운 서식 처리 (**굵게**, *기울임*, ~~취소선~~, `코드`)"""
    pattern = re.compile(r"(\*\*(.+?)\*\*|\*(.+?)\*|~~(.+?)~~|`([^`]+)`)")

    last_end = 0
    for match in pattern.finditer(text):
        # 매치 앞 텍스트
        if match.start() > last_end:
            run = paragraph.add_run(text[last_end:match.start()])
            run.font.color.rgb = _COLORS["body"]

        if match.group(2):  # **bold**
            run = paragraph.add_run(match.group(2))
            run.bold = True
            run.font.color.rgb = _COLORS["h2"]
        elif match.group(3):  # *italic*
            run = paragraph.add_run(match.group(3))
            run.italic = True
            run.font.color.rgb = _COLORS["h4"]
        elif match.group(4):  # ~~strikethrough~~
            run = paragraph.add_run(match.group(4))
            run.font.strike = True
        elif match.group(5):  # `code`
            run = paragraph.add_run(match.group(5))
            run.font.color.rgb = _COLORS["accent"]
            run.font.size = Pt(10)
            # 코드 배경 shading
            rPr = run._element.get_or_add_rPr()
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F1F5F9" w:val="clear"/>')
            rPr.append(shading)

        last_end = match.end()

    # 나머지 텍스트
    if last_end < len(text):
        run = paragraph.add_run(text[last_end:])
        run.font.color.rgb = _COLORS["body"]


def _add_table(doc: Document, table_lines: list):
    """마크다운 테이블을 웹 프리뷰 스타일과 동일한 docx 테이블로 변환"""
    if len(table_lines) < 2:
        return

    # 구분선(---|---) 행 제거
    data_lines = []
    for line in table_lines:
        cells = [c.strip() for c in line.strip("|").split("|")]
        if all(re.match(r"^[-:]+$", c) for c in cells if c):
            continue
        data_lines.append(cells)

    if not data_lines:
        return

    num_cols = max(len(row) for row in data_lines)
    num_rows = len(data_lines)

    table = doc.add_table(rows=num_rows, cols=num_cols)
    try:
        table.style = "Table Grid"
    except KeyError:
        pass  # 템플릿에 해당 스타일이 없으면 기본 스타일 사용
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for row_idx, row_data in enumerate(data_lines):
        row = table.rows[row_idx]
        for col_idx, cell_text in enumerate(row_data):
            if col_idx < num_cols:
                cell = table.cell(row_idx, col_idx)
                # 기존 텍스트 클리어 후 run으로 추가
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(cell_text.strip())
                run.font.name = "맑은 고딕"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")

                # 셀 패딩
                _set_paragraph_spacing(p, line_spacing=1.4, space_before=Pt(3), space_after=Pt(3))

                if row_idx == 0:
                    # 헤더 행: 굵게, 배경색, 텍스트 색상
                    run.bold = True
                    run.font.size = Pt(10)
                    run.font.color.rgb = _COLORS["table_header_text"]
                    _set_cell_shading(cell, "F1F5F9")
                    # 헤더 하단 보더 강조
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    borders = parse_xml(
                        f'<w:tcBorders {nsdecls("w")}>'
                        f'  <w:bottom w:val="single" w:sz="8" w:space="0" w:color="6366F1"/>'
                        f'</w:tcBorders>'
                    )
                    tcPr.append(borders)
                else:
                    # 데이터 행
                    run.font.size = Pt(10)
                    run.font.color.rgb = _COLORS["table_body_text"]
                    # 짝수 행 배경색 (교차 행, 0-based에서 row_idx 짝수 = 데이터 홀수행)
                    if row_idx % 2 == 0:
                        _set_cell_shading(cell, "F8FAFC")

    # 테이블 후 여백
    spacing_p = doc.add_paragraph("")
    spacing_p.paragraph_format.space_before = Pt(4)
    spacing_p.paragraph_format.space_after = Pt(4)
