"""OfficeMaker 설계문서 Word 파일 생성 스크립트"""
import os
import sys
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', 'server'))

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime


def set_cell_shading(cell, color):
    """셀 배경색 설정"""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_styled_table(doc, headers, rows, col_widths=None):
    """스타일이 적용된 테이블 추가"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 헤더 행
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, "2B579A")

    # 데이터 행
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(value)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F2F6FC")

    # 테이블 테두리
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(width)

    return table


def add_heading_with_color(doc, text, level, color=None):
    """색상이 적용된 헤딩 추가"""
    h = doc.add_heading(text, level=level)
    if color:
        for run in h.runs:
            run.font.color.rgb = color
    return h


def add_bullet(doc, text, level=0, bold_prefix=None):
    """불릿 포인트 추가"""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(10)
        run = p.add_run(text)
        run.font.size = Pt(10)
    else:
        run = p.add_run(text)
        run.font.size = Pt(10)
    return p


def add_body(doc, text):
    """본문 텍스트 추가"""
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(10)
    return p


def create_design_document():
    doc = Document()

    # 페이지 설정 (A4)
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # 기본 폰트 설정
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Malgun Gothic'
    font.size = Pt(10)

    # ========================================
    # 표지
    # ========================================
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("OfficeMaker")
    run.bold = True
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(0x2B, 0x57, 0x9A)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("기업용 AI 오피스 문서 자동 생성 솔루션")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()

    divider = doc.add_paragraph()
    divider.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = divider.add_run("━" * 30)
    run.font.color.rgb = RGBColor(0x2B, 0x57, 0x9A)

    doc.add_paragraph()

    doctype = doc.add_paragraph()
    doctype.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = doctype.add_run("솔루션 설계문서")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    for _ in range(4):
        doc.add_paragraph()

    info_lines = [
        f"문서 버전: 1.0",
        f"작성일: {datetime.date.today().strftime('%Y년 %m월 %d일')}",
        "분류: 내부용",
    ]
    for line in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_page_break()

    # ========================================
    # 목차
    # ========================================
    add_heading_with_color(doc, "목차", 1, RGBColor(0x2B, 0x57, 0x9A))

    toc_items = [
        "1. 솔루션 개요",
        "2. 솔루션 특징 및 차별화",
        "3. 시스템 아키텍처",
        "4. 주요 기능 상세",
        "5. 데이터 모델",
        "6. API 설계",
        "7. AI/LLM 통합",
        "8. 보안 및 인증",
        "9. 배포 및 운영",
        "10. 활용 방안",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(4)
        for run in p.runs:
            run.font.size = Pt(11)

    doc.add_page_break()

    # ========================================
    # 1. 솔루션 개요
    # ========================================
    add_heading_with_color(doc, "1. 솔루션 개요", 1, RGBColor(0x2B, 0x57, 0x9A))

    add_heading_with_color(doc, "1.1 솔루션 개념", 2, RGBColor(0x2B, 0x57, 0x9A))
    add_body(doc,
        "OfficeMaker는 기업용 AI 기반 오피스 문서 자동 생성 솔루션입니다. "
        "사용자가 보유한 리소스(문서, 이미지, 웹 검색 결과 등)를 입력하면, "
        "최신 AI 기술을 활용하여 전문적인 수준의 프레젠테이션(PPTX), 스프레드시트(XLSX), "
        "워드 문서(DOCX), HTML 리포트를 자동으로 생성합니다."
    )
    add_body(doc,
        "관리자가 기업 브랜드에 맞는 슬라이드 템플릿을 사전에 구축하고, "
        "사용자는 간단한 지침과 리소스만 등록하면 AI가 최적의 문서를 자동으로 생성합니다. "
        "OnlyOffice 통합을 통해 생성된 문서를 브라우저에서 직접 편집할 수 있으며, "
        "다수 사용자가 동시에 협업할 수 있는 실시간 공동 편집 기능을 제공합니다."
    )

    add_heading_with_color(doc, "1.2 솔루션 비전", 2, RGBColor(0x2B, 0x57, 0x9A))
    add_body(doc,
        "OfficeMaker는 기업 내 반복적인 문서 작업의 생산성을 혁신적으로 향상시키는 것을 목표로 합니다. "
        "AI가 리소스를 분석하고 기업 표준 템플릿에 맞춰 문서를 생성함으로써, "
        "사용자는 콘텐츠 기획과 의사결정에 집중할 수 있습니다."
    )

    add_styled_table(doc,
        ["구분", "내용"],
        [
            ["솔루션명", "OfficeMaker (K-Portal OfficeMaker)"],
            ["솔루션 유형", "기업용 AI 오피스 문서 자동 생성 플랫폼"],
            ["지원 문서", "PowerPoint(PPTX), Excel(XLSX), Word(DOCX), HTML Report, PDF"],
            ["AI 엔진", "Claude (Anthropic), Google Gemini, Perplexity Search"],
            ["대상 사용자", "기업 임직원 (관리자 + 일반 사용자)"],
            ["지원 언어", "한국어, 영어, 일본어, 중국어"],
            ["배포 방식", "온프레미스 / 프라이빗 클라우드"],
        ],
        col_widths=[1.5, 5.0]
    )

    doc.add_page_break()

    # ========================================
    # 2. 솔루션 특징 및 차별화
    # ========================================
    add_heading_with_color(doc, "2. 솔루션 특징 및 차별화", 1, RGBColor(0x2B, 0x57, 0x9A))

    add_heading_with_color(doc, "2.1 핵심 특징", 2, RGBColor(0x2B, 0x57, 0x9A))

    features = [
        ("AI 기반 멀티포맷 문서 생성: ",
         "하나의 리소스 풀에서 PPTX, XLSX, DOCX, HTML 리포트 등 다양한 형식의 문서를 AI가 자동으로 생성합니다."),
        ("기업 맞춤 템플릿 시스템: ",
         "관리자가 기업 CI/BI에 맞는 슬라이드 템플릿을 비주얼 캔버스에서 직접 설계하고, 사용자는 해당 템플릿 기반으로 일관된 품질의 문서를 생성합니다."),
        ("인포그래픽 자동 생성: ",
         "Google Gemini AI를 활용하여 슬라이드별 인포그래픽 배경 이미지를 자동으로 생성합니다. 스타일 힌트와 레퍼런스 이미지를 기반으로 일관된 디자인을 유지합니다."),
        ("실시간 스트리밍 생성: ",
         "SSE(Server-Sent Events) 기반으로 문서 생성 과정을 실시간으로 표시합니다. 사용자는 생성 진행 상황을 타이핑 효과로 확인할 수 있습니다."),
        ("다중 리소스 통합: ",
         "파일 업로드(Word, Excel, PPT, PDF, 이미지), 텍스트 입력, 웹 검색(Perplexity API), URL 크롤링, YouTube 자막 추출 등 다양한 소스의 리소스를 통합 관리합니다."),
        ("OnlyOffice 협업 편집: ",
         "생성된 PPTX, XLSX, DOCX 문서를 OnlyOffice Document Server와 연동하여 브라우저에서 직접 편집할 수 있습니다. 별도의 오피스 프로그램 설치가 불필요합니다."),
        ("실시간 공동 작업: ",
         "Redis 기반 슬라이드 잠금(Locking)과 온라인 프레즌스(Presence) 추적으로 다수 사용자가 동시에 문서를 편집할 수 있습니다."),
        ("다국어 지원: ",
         "한국어, 영어, 일본어, 중국어 4개 언어로 문서를 생성할 수 있으며, 기존 문서의 번역 기능도 제공합니다."),
    ]

    for prefix, text in features:
        add_bullet(doc, text, bold_prefix=prefix)

    add_heading_with_color(doc, "2.2 차별화 포인트", 2, RGBColor(0x2B, 0x57, 0x9A))

    add_styled_table(doc,
        ["차별화 요소", "일반 AI 문서 도구", "OfficeMaker"],
        [
            ["템플릿 관리", "범용 템플릿 또는 없음", "기업 맞춤 템플릿 비주얼 에디터"],
            ["조직 연동", "별도 인증", "조직도 DB 직접 연동, SSO/JWT"],
            ["문서 형식", "단일 형식", "PPTX + XLSX + DOCX + HTML + PDF"],
            ["인포그래픽", "수동 이미지 삽입", "AI 자동 이미지 생성 (Gemini)"],
            ["협업", "개인 작업", "실시간 공동 편집 + 슬라이드 잠금"],
            ["배포 방식", "SaaS (클라우드)", "온프레미스/프라이빗 클라우드 지원"],
            ["보안", "외부 클라우드", "사내 서버 운영, JWT 인증"],
            ["커스터마이징", "제한적", "프롬프트 관리, HTML 스킬, 템플릿 자유 설계"],
        ],
        col_widths=[1.5, 2.5, 2.5]
    )

    doc.add_page_break()

    # ========================================
    # 3. 시스템 아키텍처
    # ========================================
    add_heading_with_color(doc, "3. 시스템 아키텍처", 1, RGBColor(0x2B, 0x57, 0x9A))

    add_heading_with_color(doc, "3.1 기술 스택", 2, RGBColor(0x2B, 0x57, 0x9A))

    add_styled_table(doc,
        ["계층", "기술", "용도"],
        [
            ["Backend", "Python 3.x + FastAPI", "비동기 REST API 서버"],
            ["Database", "MongoDB v4 + Motor", "비동기 NoSQL 데이터 저장"],
            ["Cache", "Redis", "세션, 잠금, 프레즌스, 캐시"],
            ["Frontend", "HTML5 + CSS + JS + jQuery", "SPA 방식 웹 클라이언트"],
            ["UI Framework", "Tailwind CSS", "유틸리티 기반 스타일링"],
            ["Chart", "Chart.js", "프론트엔드 차트 렌더링"],
            ["AI - 텍스트", "Claude API (Anthropic)", "문서 콘텐츠 생성"],
            ["AI - 이미지", "Google Gemini API", "인포그래픽 이미지 생성"],
            ["AI - 검색", "Perplexity API", "웹 검색 리소스"],
            ["문서 생성", "python-pptx, openpyxl, python-docx", "서버 사이드 문서 파일 생성"],
            ["문서 편집", "OnlyOffice Document Server", "브라우저 기반 문서 편집"],
            ["인증", "JWT (HS256)", "토큰 기반 사용자 인증"],
        ],
        col_widths=[1.2, 2.8, 2.5]
    )

    add_heading_with_color(doc, "3.2 시스템 구성도", 2, RGBColor(0x2B, 0x57, 0x9A))

    add_body(doc, "OfficeMaker는 다음과 같은 계층 구조로 구성됩니다:")
    doc.add_paragraph()

    arch_text = (
        "┌─────────────────────────────────────────────────────────────────┐\n"
        "│                    클라이언트 계층 (Browser)                      │\n"
        "│  ┌──────────────┐  ┌──────────────┐  ┌───────────────────┐     │\n"
        "│  │  사용자 모듈   │  │  관리자 모듈   │  │  OnlyOffice Editor │    │\n"
        "│  │  (front/)     │  │  (admin/)     │  │  (외부 서버)        │    │\n"
        "│  └──────┬───────┘  └──────┬───────┘  └────────┬──────────┘    │\n"
        "└─────────┼─────────────────┼───────────────────┼───────────────┘\n"
        "          │ REST API / SSE  │                   │ Callback\n"
        "┌─────────┼─────────────────┼───────────────────┼───────────────┐\n"
        "│         ▼                 ▼                   ▼               │\n"
        "│              FastAPI Application Server                       │\n"
        "│  ┌──────────┐ ┌──────────┐ ┌──────────┐ ┌──────────┐        │\n"
        "│  │  Routers  │ │ Services │ │  Models  │ │  Utils   │        │\n"
        "│  └──────────┘ └──────────┘ └──────────┘ └──────────┘        │\n"
        "└──────┬──────────────┬──────────────┬─────────────────────────┘\n"
        "       │              │              │\n"
        "  ┌────▼────┐   ┌────▼────┐   ┌────▼────────────────────┐\n"
        "  │ MongoDB  │   │  Redis  │   │   External AI APIs      │\n"
        "  │ (데이터)  │   │ (캐시)  │   │  Claude / Gemini /     │\n"
        "  │          │   │         │   │  Perplexity             │\n"
        "  └─────────┘   └─────────┘   └─────────────────────────┘\n"
    )

    p = doc.add_paragraph()
    run = p.add_run(arch_text)
    run.font.size = Pt(8)
    run.font.name = 'Consolas'

    add_heading_with_color(doc, "3.3 프로젝트 디렉토리 구조", 2, RGBColor(0x2B, 0x57, 0x9A))

    dir_structure = (
        "OfficeMaker/\n"
        "├── .env                         # 환경설정 (API 키, DB 접속 정보)\n"
        "├── server/                      # FastAPI 백엔드 서버\n"
        "│   ├── main.py                  # 엔트리포인트\n"
        "│   ├── config.py                # 설정 로더\n"
        "│   ├── routers/                 # API 라우터\n"
        "│   │   ├── auth.py              # 인증 API\n"
        "│   │   ├── admin.py             # 관리자 API\n"
        "│   │   ├── template.py          # 템플릿 관리\n"
        "│   │   ├── project.py           # 프로젝트 관리\n"
        "│   │   ├── resource.py          # 리소스 관리\n"
        "│   │   ├── generate.py          # 문서 생성\n"
        "│   │   ├── prompt.py            # AI 프롬프트 관리\n"
        "│   │   ├── onlyoffice.py        # OnlyOffice 연동\n"
        "│   │   └── font.py              # 폰트 관리\n"
        "│   ├── services/                # 비즈니스 로직\n"
        "│   │   ├── mongo_service.py     # MongoDB 연결\n"
        "│   │   ├── auth_service.py      # 인증 서비스\n"
        "│   │   ├── llm_service.py       # AI 모델 연동\n"
        "│   │   ├── ppt_service.py       # PPTX 생성\n"
        "│   │   ├── excel_service.py     # XLSX 생성\n"
        "│   │   ├── word_service.py      # DOCX 생성\n"
        "│   │   ├── infographic_service.py # 인포그래픽 생성\n"
        "│   │   ├── search_service.py    # 웹 검색\n"
        "│   │   ├── onlyoffice_service.py # OnlyOffice 서비스\n"
        "│   │   └── file_service.py      # 파일 처리\n"
        "│   ├── models/                  # 데이터 모델\n"
        "│   └── utils/                   # 유틸리티\n"
        "├── admin/                       # 관리자 프론트엔드\n"
        "│   ├── index.html\n"
        "│   ├── css/admin.css\n"
        "│   └── js/admin.js\n"
        "├── front/                       # 사용자 프론트엔드\n"
        "│   ├── index.html\n"
        "│   ├── css/app.css\n"
        "│   └── js/app.js\n"
        "└── uploads/                     # 업로드 파일 저장\n"
    )

    p = doc.add_paragraph()
    run = p.add_run(dir_structure)
    run.font.size = Pt(8)
    run.font.name = 'Consolas'

    doc.add_page_break()

    # ========================================
    # 4. 주요 기능 상세
    # ========================================
    add_heading_with_color(doc, "4. 주요 기능 상세", 1, RGBColor(0x2B, 0x57, 0x9A))

    # 4.1 사용자 기능
    add_heading_with_color(doc, "4.1 사용자 기능", 2, RGBColor(0x2B, 0x57, 0x9A))

    add_heading_with_color(doc, "4.1.1 프로젝트 관리", 3, RGBColor(0x2B, 0x57, 0x9A))
    add_body(doc, "사용자는 프로젝트 단위로 문서 생성 작업을 관리합니다.")
    bullets = [
        ("프로젝트 유형: ", "슬라이드(PPTX), 엑셀(XLSX), 워드(DOCX), OnlyOffice PPTX/XLSX/DOCX"),
        ("폴더 관리: ", "프로젝트를 폴더별로 분류하여 체계적으로 관리"),
        ("상태 추적: ", "초안(draft) → 준비(preparing) → 생성중(generating) → 완료(generated)"),
        ("공유: ", "다른 사용자에게 편집자(editor) 또는 뷰어(viewer) 권한으로 프로젝트 공유"),
        ("검색 및 필터: ", "프로젝트명 검색, 폴더별 필터링, 최근 사용 순 정렬"),
    ]
    for prefix, text in bullets:
        add_bullet(doc, text, bold_prefix=prefix)

    add_heading_with_color(doc, "4.1.2 리소스 관리", 3, RGBColor(0x2B, 0x57, 0x9A))
    add_body(doc, "AI 문서 생성의 입력 소스가 되는 다양한 형태의 리소스를 등록하고 관리합니다.")

    add_styled_table(doc,
        ["리소스 유형", "지원 형식", "처리 방식"],
        [
            ["파일 업로드", "DOCX, XLSX, PPTX, PDF, TXT, CSV, MD", "텍스트 자동 추출 + 캐싱(SHA-256)"],
            ["이미지 업로드", "PNG, JPG, GIF 등", "Claude Vision API로 이미지 분석"],
            ["텍스트 입력", "자유 형식", "직접 붙여넣기"],
            ["웹 검색", "Perplexity API", "키워드 검색 → URL + 콘텐츠 추출"],
            ["URL 입력", "웹 페이지, YouTube", "웹 스크래핑, YouTube 자막 추출"],
        ],
        col_widths=[1.3, 2.5, 2.7]
    )

    add_heading_with_color(doc, "4.1.3 문서 생성", 3, RGBColor(0x2B, 0x57, 0x9A))
    add_body(doc, "등록된 리소스와 사용자 지침을 기반으로 AI가 문서를 자동 생성합니다.")

    gen_bullets = [
        ("슬라이드 생성: ", "템플릿 선택 → 리소스 분석 → 슬라이드별 콘텐츠 매핑 → 실시간 스트리밍 표시"),
        ("인포그래픽 모드: ", "Gemini AI로 슬라이드별 배경 이미지 자동 생성, 커버/콘텐츠 슬라이드 구분"),
        ("엑셀 생성: ", "리소스에서 데이터를 추출하여 시트, 컬럼, 차트를 자동 구성"),
        ("워드 생성: ", "Markdown 기반 콘텐츠 생성 후 DOCX 변환, 표/차트 자동 삽입"),
        ("HTML 리포트: ", "관리자가 등록한 HTML 스킬 기반으로 맞춤형 리포트 생성"),
        ("번역: ", "생성된 문서 전체를 다른 언어로 자동 번역"),
    ]
    for prefix, text in gen_bullets:
        add_bullet(doc, text, bold_prefix=prefix)

    add_heading_with_color(doc, "4.1.4 문서 편집 및 미리보기", 3, RGBColor(0x2B, 0x57, 0x9A))
    edit_bullets = [
        ("슬라이드 캔버스 편집: ", "오브젝트 선택, 드래그 이동, 리사이즈, 텍스트 편집, 이미지 교체"),
        ("AI 슬라이드 디자인: ", "개별 슬라이드에 대해 AI가 자동으로 레이아웃과 디자인을 최적화"),
        ("OnlyOffice 편집: ", "생성된 PPTX/XLSX/DOCX를 브라우저에서 네이티브 편집"),
        ("프레젠테이션 모드: ", "풀스크린 프레젠테이션 미리보기 지원"),
        ("공유 링크: ", "만료일이 설정된 공유 URL을 생성하여 외부 공유"),
    ]
    for prefix, text in edit_bullets:
        add_bullet(doc, text, bold_prefix=prefix)

    add_heading_with_color(doc, "4.1.5 다운로드", 3, RGBColor(0x2B, 0x57, 0x9A))
    add_body(doc, "PPTX, XLSX, DOCX 원본 파일 다운로드 및 PDF 변환 다운로드를 지원합니다. "
             "슬라이드의 경우 인포그래픽 이미지, 차트, 표 등 모든 오브젝트가 포함된 완전한 파일을 생성합니다.")

    # 4.2 관리자 기능
    add_heading_with_color(doc, "4.2 관리자 기능", 2, RGBColor(0x2B, 0x57, 0x9A))

    add_heading_with_color(doc, "4.2.1 템플릿 관리", 3, RGBColor(0x2B, 0x57, 0x9A))
    add_body(doc, "관리자는 비주얼 캔버스 에디터를 통해 슬라이드 템플릿을 설계하고 관리합니다.")
    tmpl_bullets = [
        ("비주얼 캔버스: ", "960px 기반 캔버스에서 드래그&드롭으로 오브젝트 배치"),
        ("오브젝트 유형: ", "텍스트, 이미지, 도형(사각형, 원형, 선, 화살표), 표, 차트"),
        ("오브젝트 속성: ", "위치, 크기, 색상, 폰트, 정렬, z-index 등 세밀한 속성 제어"),
        ("차트 유형: ", "Bar, Line, Pie, Doughnut, Area, Scatter, Radar"),
        ("슬라이드 메타: ", "제목 유무, 거버넌스, 항목 수 등 AI가 활용할 메타정보 설정"),
        ("발행 관리: ", "템플릿 발행/비발행 상태 관리"),
        ("배경 이미지: ", "공통 배경이미지 업로드 및 관리"),
        ("슬라이드 사이즈: ", "16:9, 4:3, A4 비율 선택"),
    ]
    for prefix, text in tmpl_bullets:
        add_bullet(doc, text, bold_prefix=prefix)

    add_heading_with_color(doc, "4.2.2 AI 프롬프트 관리", 3, RGBColor(0x2B, 0x57, 0x9A))
    add_body(doc, "AI 문서 생성에 사용되는 시스템 프롬프트와 사용자 프롬프트를 데이터베이스에서 관리합니다. "
             "프롬프트별로 사용할 AI 모델(Opus, Sonnet, Haiku)을 지정할 수 있으며, "
             "변수 바인딩({lang_instruction}, {resources_text} 등)을 통해 동적 프롬프트를 구성합니다.")

    add_heading_with_color(doc, "4.2.3 폰트 관리", 3, RGBColor(0x2B, 0x57, 0x9A))
    add_body(doc, "시스템에서 사용할 폰트를 등록하고 관리합니다. 웹 폰트 URL을 등록하면 "
             "프론트엔드에서 동적으로 로딩됩니다. 전체 템플릿에 대한 폰트 일괄 변경 기능도 제공합니다.")

    add_heading_with_color(doc, "4.2.4 HTML 스킬 관리", 3, RGBColor(0x2B, 0x57, 0x9A))
    add_body(doc, "맞춤형 HTML 리포트 템플릿(스킬)을 생성하고 관리합니다. "
             "HTML 템플릿과 CSS를 직접 편집하여 기업 특화 리포트 형식을 정의할 수 있습니다.")

    doc.add_page_break()

    # 4.3 협업 기능
    add_heading_with_color(doc, "4.3 협업 기능", 2, RGBColor(0x2B, 0x57, 0x9A))

    add_styled_table(doc,
        ["기능", "설명", "기술 구현"],
        [
            ["프로젝트 공유", "Owner/Editor/Viewer 역할 기반 접근 제어", "MongoDB collaborators 컬렉션"],
            ["슬라이드 잠금", "편집 중인 슬라이드 자동 잠금 (5분 TTL)", "Redis 기반 분산 잠금"],
            ["온라인 프레즌스", "현재 접속 중인 사용자 표시", "Redis + 하트비트"],
            ["편집 이력", "슬라이드별 변경 이력 기록", "MongoDB slide_history"],
            ["OnlyOffice 공동편집", "다수 사용자 동시 문서 편집", "OnlyOffice Document Server"],
        ],
        col_widths=[1.5, 2.5, 2.5]
    )

    # 4.4 외부 API
    add_heading_with_color(doc, "4.4 외부 연동 API", 2, RGBColor(0x2B, 0x57, 0x9A))
    add_body(doc, "외부 시스템에서 OfficeMaker의 문서 생성 기능을 RESTful API로 호출할 수 있습니다. "
             "파일과 지침을 전송하면 자동으로 슬라이드를 생성하고 공유 URL을 반환합니다. "
             "인포그래픽 모드, 출력 언어 선택 등의 옵션을 지원합니다.")

    doc.add_page_break()

    # ========================================
    # 5. 데이터 모델
    # ========================================
    add_heading_with_color(doc, "5. 데이터 모델", 1, RGBColor(0x2B, 0x57, 0x9A))

    add_heading_with_color(doc, "5.1 MongoDB 컬렉션", 2, RGBColor(0x2B, 0x57, 0x9A))
    add_body(doc, "OfficeMaker는 MongoDB를 주 데이터 저장소로 사용하며, 다음과 같은 컬렉션으로 구성됩니다.")

    add_styled_table(doc,
        ["컬렉션", "용도", "주요 필드"],
        [
            ["projects", "프로젝트 관리", "user_key, name, project_type, status, template_id"],
            ["resources", "입력 리소스", "project_id, resource_type, title, content, file_path"],
            ["generated_slides", "생성된 슬라이드", "project_id, template_slide_id, objects, order"],
            ["generated_excel", "생성된 엑셀", "project_id, meta, sheets(columns, rows, charts)"],
            ["generated_docx", "생성된 워드", "project_id, meta, sections(level, title, content)"],
            ["templates", "슬라이드 템플릿", "name, is_published, slide_size, background_image"],
            ["slides", "템플릿 슬라이드", "template_id, order, objects, slide_meta"],
            ["custom_templates", "사용자 PPTX 템플릿", "project_id, file_path, slides_count"],
            ["collaborators", "프로젝트 공유", "project_id, user_key, role, invited_by"],
            ["onlyoffice_documents", "OnlyOffice 파일", "project_id, file_type, document_key"],
            ["prompts", "AI 프롬프트", "key, name, model, content"],
            ["html_skills", "HTML 스킬", "name, html_template, css_template"],
            ["fonts", "등록 폰트", "name, family, url"],
            ["project_folders", "프로젝트 폴더", "user_key, name, order"],
            ["file_text_cache", "파일 텍스트 캐시", "file_hash(SHA-256), content"],
        ],
        col_widths=[1.5, 1.5, 3.5]
    )

    add_heading_with_color(doc, "5.2 외부 데이터베이스", 2, RGBColor(0x2B, 0x57, 0x9A))
    add_body(doc, "조직도 정보는 별도의 MongoDB 데이터베이스(im_org_info)의 user_info 컬렉션에서 조회합니다.")

    add_styled_table(doc,
        ["필드", "설명", "예시"],
        [
            ["ky", "사용자 고유 키", "user001"],
            ["nm", "사용자명", "홍길동"],
            ["dp", "부서명", "개발팀"],
            ["em", "이메일", "hong@company.com"],
            ["role", "역할", "admin / (없음)"],
            ["m365", "M365 계정", "hong@company.onmicrosoft.com"],
        ],
        col_widths=[1.0, 2.0, 3.5]
    )

    add_heading_with_color(doc, "5.3 캐시 전략 (Redis)", 2, RGBColor(0x2B, 0x57, 0x9A))

    add_styled_table(doc,
        ["키 패턴", "용도", "TTL"],
        [
            ["fonts:all", "전체 폰트 목록", "7일"],
            ["fonts:public", "공개 폰트 목록", "7일"],
            ["templates:{id}", "템플릿 데이터", "7일"],
            ["officemaker:{id}:locks", "슬라이드 잠금", "5분"],
            ["officemaker:{id}:presence", "온라인 사용자", "5분"],
            ["officemaker:{id}:cancel", "생성 취소 플래그", "가변"],
        ],
        col_widths=[2.5, 2.0, 2.0]
    )

    doc.add_page_break()

    # ========================================
    # 6. API 설계
    # ========================================
    add_heading_with_color(doc, "6. API 설계", 1, RGBColor(0x2B, 0x57, 0x9A))

    add_heading_with_color(doc, "6.1 API 라우팅 규칙", 2, RGBColor(0x2B, 0x57, 0x9A))
    add_body(doc, "OfficeMaker는 JWT 토큰을 URL 경로에 포함하는 방식으로 인증을 처리합니다.")

    add_styled_table(doc,
        ["패턴", "설명", "인증"],
        [
            ["/{jwt}/api/...", "사용자 API", "JWT 필수"],
            ["/{jwt}/api/admin/...", "관리자 API", "JWT + admin 역할"],
            ["/api/auth/...", "인증 API", "불필요"],
            ["/api/shared/{token}/...", "공유 API", "공유 토큰"],
            ["/api/onlyoffice/...", "OnlyOffice 콜백", "불필요"],
            ["/api/external/...", "외부 연동 API", "API 키"],
        ],
        col_widths=[2.0, 2.0, 2.5]
    )

    add_heading_with_color(doc, "6.2 주요 API 엔드포인트", 2, RGBColor(0x2B, 0x57, 0x9A))

    # 프로젝트 API
    add_heading_with_color(doc, "프로젝트 관리 API", 3, RGBColor(0x2B, 0x57, 0x9A))
    add_styled_table(doc,
        ["Method", "Endpoint", "설명"],
        [
            ["GET", "/{jwt}/api/projects", "프로젝트 목록 조회"],
            ["POST", "/{jwt}/api/projects", "프로젝트 생성"],
            ["GET", "/{jwt}/api/projects/{id}", "프로젝트 상세 조회"],
            ["PUT", "/{jwt}/api/projects/{id}", "프로젝트 수정"],
            ["DELETE", "/{jwt}/api/projects/{id}", "프로젝트 삭제"],
            ["POST", "/{jwt}/api/projects/{id}/reset", "생성 콘텐츠 초기화"],
        ],
        col_widths=[0.8, 2.7, 3.0]
    )

    # 리소스 API
    add_heading_with_color(doc, "리소스 관리 API", 3, RGBColor(0x2B, 0x57, 0x9A))
    add_styled_table(doc,
        ["Method", "Endpoint", "설명"],
        [
            ["GET", "/{jwt}/api/resources/{project_id}", "리소스 목록 조회"],
            ["POST", "/{jwt}/api/resources/text", "텍스트 리소스 등록"],
            ["POST", "/{jwt}/api/resources/file", "파일 업로드 (자동 텍스트 추출)"],
            ["POST", "/{jwt}/api/resources/image", "이미지 업로드 (AI 분석)"],
            ["POST", "/{jwt}/api/resources/web-search", "웹 검색 리소스"],
            ["POST", "/{jwt}/api/resources/url", "URL 리소스 등록"],
        ],
        col_widths=[0.8, 2.7, 3.0]
    )

    # 생성 API
    add_heading_with_color(doc, "문서 생성 API", 3, RGBColor(0x2B, 0x57, 0x9A))
    add_styled_table(doc,
        ["Method", "Endpoint", "설명"],
        [
            ["POST", "/{jwt}/api/generate/stream", "슬라이드 스트리밍 생성"],
            ["POST", "/{jwt}/api/generate/infographic-batch", "인포그래픽 일괄 생성"],
            ["POST", "/{jwt}/api/generate/excel/stream", "엑셀 스트리밍 생성"],
            ["POST", "/{jwt}/api/generate/docx/stream", "워드 스트리밍 생성"],
            ["POST", "/{jwt}/api/generate/html-report", "HTML 리포트 생성"],
            ["POST", "/{jwt}/api/generate/translate", "문서 번역"],
            ["POST", "/{jwt}/api/generate/ai-slide", "AI 슬라이드 디자인"],
            ["GET", "/{jwt}/api/projects/{id}/download/pptx", "PPTX 다운로드"],
            ["GET", "/{jwt}/api/projects/{id}/download/pdf", "PDF 다운로드"],
            ["GET", "/{jwt}/api/projects/{id}/download/xlsx", "XLSX 다운로드"],
        ],
        col_widths=[0.8, 2.7, 3.0]
    )

    doc.add_page_break()

    # ========================================
    # 7. AI/LLM 통합
    # ========================================
    add_heading_with_color(doc, "7. AI/LLM 통합", 1, RGBColor(0x2B, 0x57, 0x9A))

    add_heading_with_color(doc, "7.1 AI 모델 구성", 2, RGBColor(0x2B, 0x57, 0x9A))

    add_styled_table(doc,
        ["AI 서비스", "모델", "용도", "특징"],
        [
            ["Claude (Anthropic)", "Opus 4.6", "슬라이드 콘텐츠 생성 (고정밀)", "최고 품질, 복잡한 구조화"],
            ["Claude (Anthropic)", "Sonnet 4.6", "아웃라인, 엑셀, 워드 생성", "고성능 대비 빠른 속도"],
            ["Claude (Anthropic)", "Haiku 4.5", "빠른 작업 처리", "경량 태스크"],
            ["Google Gemini", "gemini-3.1-flash", "인포그래픽 이미지 생성", "이미지 생성 특화"],
            ["Perplexity", "검색 API", "웹 검색 리소스", "실시간 웹 검색"],
        ],
        col_widths=[1.3, 1.3, 2.0, 2.0]
    )

    add_heading_with_color(doc, "7.2 프롬프트 관리 체계", 2, RGBColor(0x2B, 0x57, 0x9A))
    add_body(doc, "AI 프롬프트는 데이터베이스(prompts 컬렉션)에 저장되어 관리자가 웹 UI에서 직접 수정할 수 있습니다. "
             "이를 통해 서버 재배포 없이 AI의 동작을 조정할 수 있습니다.")

    prompt_bullets = [
        ("시스템 프롬프트: ", "AI의 역할과 출력 형식을 정의 (slide_generation_system, excel_generation_system 등)"),
        ("사용자 프롬프트: ", "리소스, 지침, 템플릿 정보를 변수로 바인딩 ({resources_text}, {instructions} 등)"),
        ("모델 선택: ", "프롬프트별로 사용할 AI 모델을 개별 지정 가능"),
        ("인포그래픽 프롬프트: ", "커버/콘텐츠 이미지 생성, 스타일 오버라이드, 레퍼런스 이미지 프롬프트"),
    ]
    for prefix, text in prompt_bullets:
        add_bullet(doc, text, bold_prefix=prefix)

    add_heading_with_color(doc, "7.3 API 키 관리", 2, RGBColor(0x2B, 0x57, 0x9A))
    add_body(doc, "복수의 API 키를 환경변수에 쉼표로 구분하여 등록하면, 라운드 로빈 방식으로 키를 순환하여 "
             "API 호출 부하를 분산합니다. 이를 통해 API Rate Limit 제한을 효과적으로 관리합니다.")

    add_heading_with_color(doc, "7.4 문서 생성 흐름", 2, RGBColor(0x2B, 0x57, 0x9A))

    flow_text = (
        "1. 리소스 수집\n"
        "   사용자가 등록한 리소스(파일, 텍스트, 웹검색 등)를 수집하고 텍스트를 추출합니다.\n\n"
        "2. 템플릿 매칭\n"
        "   선택된 템플릿의 슬라이드 구조와 메타정보를 분석합니다.\n\n"
        "3. AI 콘텐츠 생성\n"
        "   시스템 프롬프트 + 리소스 + 지침 + 템플릿 정보를 Claude API에 전송합니다.\n"
        "   SSE 스트리밍으로 실시간 응답을 받습니다.\n\n"
        "4. 구조화 파싱\n"
        "   AI 응답(JSON)을 파싱하여 슬라이드별 오브젝트로 매핑합니다.\n\n"
        "5. 인포그래픽 생성 (선택)\n"
        "   인포그래픽 모드가 활성화된 경우, Gemini API로 슬라이드별 배경 이미지를 생성합니다.\n\n"
        "6. 저장 및 표시\n"
        "   생성된 콘텐츠를 MongoDB에 저장하고 프론트엔드에 실시간 표시합니다."
    )
    add_body(doc, flow_text)

    doc.add_page_break()

    # ========================================
    # 8. 보안 및 인증
    # ========================================
    add_heading_with_color(doc, "8. 보안 및 인증", 1, RGBColor(0x2B, 0x57, 0x9A))

    add_heading_with_color(doc, "8.1 인증 체계", 2, RGBColor(0x2B, 0x57, 0x9A))

    auth_bullets = [
        ("JWT 토큰: ", "HS256 알고리즘으로 서명된 JWT를 URL 경로에 포함하여 모든 API 요청을 인증합니다."),
        ("토큰 만료: ", "JWT_EXPIRE_HOURS 환경변수로 토큰 유효기간을 설정합니다 (기본 24시간)."),
        ("외부 JWT 지원: ", "외부 시스템의 JWT를 수용하여 SSO(Single Sign-On)를 지원합니다."),
        ("관리자 판별: ", "조직도 DB의 role 필드가 'admin'인 경우 관리자 권한을 부여합니다."),
        ("데모 계정: ", "데모 환경을 위한 별도 계정 체계를 지원합니다."),
    ]
    for prefix, text in auth_bullets:
        add_bullet(doc, text, bold_prefix=prefix)

    add_heading_with_color(doc, "8.2 보안 정책", 2, RGBColor(0x2B, 0x57, 0x9A))

    sec_bullets = [
        ("환경변수 관리: ", "모든 민감 정보(API 키, DB 접속 정보, 비밀키 등)는 .env 파일에서 관리하며 하드코딩을 금지합니다."),
        ("패스워드 암호화: ", "사용자 패스워드는 해시 처리하여 저장합니다."),
        ("CORS 제어: ", "허용된 도메인만 API 접근을 허용합니다."),
        ("파일 업로드 제한: ", "최대 업로드 크기를 환경변수로 제한합니다 (기본 50MB)."),
        ("SSL/TLS: ", "HTTPS 통신을 위한 인증서를 지원합니다."),
        ("OnlyOffice JWT: ", "OnlyOffice Document Server와의 통신에 별도 JWT 서명을 적용합니다."),
    ]
    for prefix, text in sec_bullets:
        add_bullet(doc, text, bold_prefix=prefix)

    add_heading_with_color(doc, "8.3 L4 이중화 대응", 2, RGBColor(0x2B, 0x57, 0x9A))
    add_body(doc, "OfficeMaker는 다중 서버 환경(L4 로드밸런서)을 고려하여 설계되었습니다.")
    l4_bullets = [
        "서버 메모리/브라우저 localStorage에 상태를 저장하지 않습니다.",
        "모든 상태는 MongoDB 또는 Redis에 저장하여 서버 간 공유합니다.",
        "Redis 기반 분산 잠금으로 동시 편집 충돌을 방지합니다.",
        "세션리스(Stateless) API 설계로 어느 서버에서든 요청을 처리할 수 있습니다.",
    ]
    for text in l4_bullets:
        add_bullet(doc, text)

    doc.add_page_break()

    # ========================================
    # 9. 배포 및 운영
    # ========================================
    add_heading_with_color(doc, "9. 배포 및 운영", 1, RGBColor(0x2B, 0x57, 0x9A))

    add_heading_with_color(doc, "9.1 시스템 요구사항", 2, RGBColor(0x2B, 0x57, 0x9A))

    add_styled_table(doc,
        ["구성요소", "요구사항"],
        [
            ["OS", "Linux / Windows Server"],
            ["Python", "3.10 이상"],
            ["MongoDB", "v4 이상"],
            ["Redis", "6.x 이상"],
            ["OnlyOffice", "Document Server 7.x 이상 (선택)"],
            ["네트워크", "외부 AI API 접근 (Claude, Gemini, Perplexity)"],
            ["저장소", "업로드 파일 + 생성 파일 저장 공간"],
        ],
        col_widths=[1.5, 5.0]
    )

    add_heading_with_color(doc, "9.2 환경설정", 2, RGBColor(0x2B, 0x57, 0x9A))
    add_body(doc, "모든 설정은 프로젝트 루트의 .env 파일에서 관리합니다. 주요 설정 항목:")

    add_styled_table(doc,
        ["설정 항목", "설명", "예시"],
        [
            ["SOLUTION_NAME", "솔루션 표시 이름", "OfficeMaker"],
            ["MONGO_URI", "MongoDB 접속 URI", "mongodb://host:port/"],
            ["PPTMAKER_DB", "메인 데이터베이스명", "PPTMaker"],
            ["JWT_SECRET", "JWT 서명 비밀키", "(암호화 문자열)"],
            ["ANTHROPIC_API_KEY", "Claude API 키 (복수 가능)", "sk-ant-..."],
            ["GOOGLE_API_KEY", "Gemini API 키 (복수 가능)", "AIzaSy..."],
            ["PERPLEXITY_API_KEY", "Perplexity API 키", "pplx-..."],
            ["REDIS_HOST / PORT", "Redis 서버 정보", "127.0.0.1:6379"],
            ["ONLYOFFICE_URL", "OnlyOffice 서버 URL", "https://docs.example.com"],
            ["SUPPORTED_LANGS", "지원 언어 목록", "ko,en,ja,zh"],
            ["SERVER_BASE_URL", "서버 외부 접근 URL", "https://example.com:5030"],
        ],
        col_widths=[1.5, 2.0, 3.0]
    )

    add_heading_with_color(doc, "9.3 실행 방법", 2, RGBColor(0x2B, 0x57, 0x9A))

    run_text = (
        "# 의존성 설치\n"
        "cd server\n"
        "python -m venv venv\n"
        "source venv/Scripts/activate    # Windows\n"
        "# source venv/bin/activate      # Linux\n"
        "pip install -r requirements.txt\n\n"
        "# 서버 실행\n"
        "python main.py\n"
        "# 또는\n"
        "uvicorn main:app --host 0.0.0.0 --port 5030 --reload"
    )
    p = doc.add_paragraph()
    run = p.add_run(run_text)
    run.font.size = Pt(9)
    run.font.name = 'Consolas'

    add_heading_with_color(doc, "9.4 접속 URL", 2, RGBColor(0x2B, 0x57, 0x9A))

    add_styled_table(doc,
        ["URL", "설명"],
        [
            ["http://host:port/main", "랜딩 페이지"],
            ["http://host:port/app", "사용자 애플리케이션"],
            ["http://host:port/admin", "관리자 페이지"],
            ["http://host:port/{jwt}", "JWT 인증 사용자 페이지"],
            ["http://host:port/shared/{token}", "공유 프레젠테이션 뷰"],
        ],
        col_widths=[3.0, 3.5]
    )

    doc.add_page_break()

    # ========================================
    # 10. 활용 방안
    # ========================================
    add_heading_with_color(doc, "10. 활용 방안", 1, RGBColor(0x2B, 0x57, 0x9A))

    add_heading_with_color(doc, "10.1 기업 내부 활용", 2, RGBColor(0x2B, 0x57, 0x9A))

    use_cases_internal = [
        ("경영 보고서 자동화: ",
         "월간/분기 실적 데이터를 리소스로 등록하면 경영진 보고용 프레젠테이션을 자동 생성합니다. "
         "기업 CI/BI가 적용된 템플릿으로 일관된 품질의 보고서를 빠르게 작성할 수 있습니다."),
        ("제안서/기획서 작성: ",
         "프로젝트 기획 문서, 기술 제안서를 AI가 자동으로 구조화하여 생성합니다. "
         "사용자는 핵심 내용만 입력하면 전문적인 문서로 변환됩니다."),
        ("교육 자료 제작: ",
         "사내 교육 콘텐츠를 리소스로 등록하면 교육용 슬라이드를 자동 생성합니다. "
         "인포그래픽 모드를 활용하면 시각적으로 풍부한 교육 자료를 만들 수 있습니다."),
        ("데이터 분석 리포트: ",
         "Excel 생성 기능으로 데이터를 정리하고 차트를 자동으로 생성합니다. "
         "Word 문서와 결합하여 종합적인 분석 리포트를 작성할 수 있습니다."),
        ("회의록/업무 보고: ",
         "회의 내용이나 업무 현황을 텍스트로 입력하면 정돈된 워드 문서로 자동 변환됩니다."),
    ]
    for prefix, text in use_cases_internal:
        add_bullet(doc, text, bold_prefix=prefix)

    add_heading_with_color(doc, "10.2 부서별 활용 시나리오", 2, RGBColor(0x2B, 0x57, 0x9A))

    add_styled_table(doc,
        ["부서", "활용 시나리오", "주요 기능"],
        [
            ["경영기획", "경영 실적 보고서, 전략 기획서", "슬라이드 + 차트 자동생성"],
            ["영업/마케팅", "제품 소개서, 영업 제안서, 시장 분석", "웹 검색 + 인포그래픽"],
            ["인사/교육", "사내 교육 자료, 인사 현황 보고", "다국어 지원 + 번역"],
            ["연구개발", "기술 보고서, 연구 결과 발표", "파일 업로드 + 데이터 분석"],
            ["재무/회계", "재무제표 분석, 예산 보고서", "엑셀 생성 + 차트"],
            ["IT/시스템", "시스템 설계서, 프로젝트 현황", "워드 + 슬라이드 통합"],
        ],
        col_widths=[1.3, 2.5, 2.7]
    )

    add_heading_with_color(doc, "10.3 외부 시스템 연동 활용", 2, RGBColor(0x2B, 0x57, 0x9A))

    ext_bullets = [
        ("그룹웨어 연동: ",
         "External API를 통해 그룹웨어, ERP 등 기존 업무 시스템에서 직접 문서 생성을 요청할 수 있습니다. "
         "업무 흐름 중에 자연스럽게 문서가 생성됩니다."),
        ("포털 임베딩: ",
         "사내 포털 페이지에 OfficeMaker를 iframe으로 임베딩하여 별도 접속 없이 사용할 수 있습니다. "
         "외부 JWT 지원으로 SSO 연동이 가능합니다."),
        ("자동화 파이프라인: ",
         "정기 보고서 등 반복 작업을 API 호출로 자동화할 수 있습니다. "
         "데이터 수집 → 문서 생성 → 공유 링크 발송까지의 프로세스를 무인 운영합니다."),
    ]
    for prefix, text in ext_bullets:
        add_bullet(doc, text, bold_prefix=prefix)

    add_heading_with_color(doc, "10.4 기대 효과", 2, RGBColor(0x2B, 0x57, 0x9A))

    add_styled_table(doc,
        ["항목", "기대 효과"],
        [
            ["문서 작성 시간 단축", "기존 수 시간 소요 → 수 분 내 초안 완성 (90% 이상 시간 절감)"],
            ["문서 품질 일관성", "기업 표준 템플릿 기반으로 일관된 디자인과 구조 유지"],
            ["업무 생산성 향상", "반복적 문서 작업에서 해방, 핵심 업무에 집중"],
            ["협업 효율화", "실시간 공동 편집으로 문서 작업의 병목 해소"],
            ["다국어 대응", "글로벌 커뮤니케이션을 위한 다국어 문서 즉시 생성"],
            ["보안 강화", "온프레미스 배포로 민감 데이터의 외부 유출 방지"],
        ],
        col_widths=[1.8, 4.7]
    )

    # 문서 저장
    output_path = os.path.join(os.path.dirname(__file__), "OfficeMaker_설계문서.docx")
    doc.save(output_path)
    print(f"설계문서 생성 완료: {output_path}")
    return output_path


if __name__ == "__main__":
    create_design_document()
